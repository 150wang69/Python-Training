# 1.编程完成以下任务：
# （1）自己创建一个20页以上的文档，粘贴文字，可以是自己喜欢的小说或者其他任意文字；
# （2）使用Document包读取word文档；
# （3）统计文档中段落的个数；
# （4）使用字典word_Count统计文档中每个字出现的次数，不统计出现在文件stopwords中的字和标点符号；

"""
数据223-王洋-22
"""

from docx import Document
import jieba
from collections import Counter
import openpyxl
#使用jieba库进行中文分词，使用collections库中的Counter类统计词频

if __name__ == '__main__':

    #第一步，读取停用词有哪些，为后续的判断做准备

    StopName = open("stopwords.txt",'r',encoding='utf-8') #打开停用词的文档
    Stop_Words = StopName.read() #读取停用词
    # print(Stop_Words) #将停用词打印出来检查是否读取

    #第二步，读取word文档，将其保存在txt文档中

    Doc = Document("朱自清散文选集.docx") #创建对象
    #将doc文件转为txt文件
    with open("朱自清散文选集.txt",'w',encoding="utf-8") as Word_File:
        for para in Doc.paragraphs: #使用paragraphs读取docx文件段落
            # print(para.text) #打印检验
            Word_File.write(para.text + "\n") #将段落存储分行

    #第三步，切割txt文件的内容 f = open(Word_File,'r',encoding='utf-8')

    txt = open("朱自清散文选集.txt","r", encoding='utf-8').read() #读取文件内容
    words = jieba.lcut(txt) #使用jieba库的lcut方法进入精确模式划分，如果想要使用全模式就加上参数cut_all = True(默认为False)
    #并且自动返回一个列表类型
    print("-----划分后的词组为：-----") #打印检查
    print(words)

    #第四步，遍历停用词表，判断是否有停用词，若有则删除。

    li = [] #使用一个新列表保存去掉停用词的内容
    #可以使用for循环也可以使用filter函数
    # 1.filter函数
    print("-----方法一：使用filter函数，去掉停用词后的词组为：-----")
    lst = filter(lambda x:x not in Stop_Words,words) #返回值为迭代器
    lt = list(lst) #得到一个可迭代对象
    #filter函数有两个参数，function和可迭代对象，直接使用lambda函数。
    print(list(lst))
    # 2.for循环
    print("-----方法二：使用for循环，去掉停用词后的词组为：-----")
    for word in words:
        if word not in Stop_Words:
            li.append(word)
    print(list(li))

    #第五步，统计词频

    counter = Counter(li) #创建Counter对象进行词频的统计，Counter的对象是一个dic的子类。
    phrase = []
    for word,num in counter.items(): #调用dic的items方法
        phrase.append((word,num)) #以元组类型添加到列表当中去
    print("-----统计的词频为：-----")
    print(phrase)

    #第六步，将统计的词频保存到Excel表格当中
    wb = openpyxl.Workbook() #创建工作薄
    ws = wb.active #由于在创建工作簿的同时也创建一个工作表，所以使用active直接调用该工作表
    #将出现的词语和词频写入表格
    #定义表头
    ws['A1'] = "词组"
    ws['B1'] = "词频"
    #循环进行索引添加
    for i, row in enumerate(phrase):#i为循环的索引，row对应phrase列表里面的索引位置的值
        #由于i是从0开始的，所以需要将i+2作为行数
        ws['A{}'.format(i+2)] = row[0] #表示会将row中的第一个元素放入工作表的A列
        ws['B{}'.format(i+2)] = row[1] #将row中的第二个元素放入工作表的B列
    wb.save("词频统计表.xlsx") #储存
    wb.close()

    Ten_max = counter.most_common(10) #Counter对象的most_common方法可以直接调用，统计词频最高的十个词语并输出
    """关于该方法的解释：
    List the n most common elements and their counts from the most common to the least. 
    If n is None, then list all element counts.
    """
    print(Ten_max) #输出结果

#获取停用词
# def stop_del(text):
#     stop_f = open(text,"r",encoding='utf-8')
#     stop_words = list()
#     for line in stop_f.readlines():
#         line = line.strip()
#         if not len(line):
#             continue
#         stop_words.append(line)
#     # stop_f.close()
#     # print(len(stop_words))
#     return stop_words
#读取停用词文件
# def readStopWords():
#     with open("stopwords.txt","r",encoding='utf-8') as file_stop:
#         stop_words = file_stop.read()


# #读取word文件
# def read_Word(txt_):
#     doc = Document(txt_)
#     with open("TEXT.txt",'w',encoding='utf-8') as file_word:
#         for par in doc.paragraphs:
#             print(par.text) #输出内容
#             text.append(par.text) #将文本加入到文件TEXT.txt当中

    # #获取文档对象
    # file = Document(txt_)
    # text = []
    # #若要在文档中写入内容，使用add_paragraph()方法
    # #file.add_paragraph("...")
    #
    # #获取总段落数
    # print("段落数有" + str(len(file.paragraphs))) #每个回车隔离一行，共105行
    # #输出每一段的内容，同时将内容存入列表中
    #
    #
    #      #添加
    # return ' '.join(text)

#分词并统计词频
# def count_words(text):
#     lst = [] #设置一个空列表用来储存切割后的字符
    # with open("Text.txt",)
    # cutting = list()
    # with open("str_all.txt",'w') as f:
    #     f.write(text)
    #     f.close()
    # file = open("str_all.txt","r",encoding='utf-8')
    # for line in file.readlines():
    #     line = line.strip()
    #     if not len(line):
    #         continue
    #     outstr = ''
    #     seg_list = jieba.cut(line, cut_all=False)
    #     for word in seg_list:
    #         if word not in stop_del(stop_words_file):
    #             if word != '\t':
    #                 outstr += word
    #                 outstr += " "
    #                 # seg_list = " ".join(seg_list)
    #     cutting.append(outstr.strip())
    # file.close()
    # words = jieba.lcut(text)  #精简模式，返回一个列表类型的结果
    # word_counts =Counter(words)
    # return word_counts

#将结果输出道EXCEl上面
# def outputToEXCEL(word_counts):
#     keywords = []
#     for word , count in word_counts.items():
#         keywords.append([word,count])
#
#     wb = openpyxl.Workbook() #先创建工作簿
#     ws = wb.active
#
#     #将出现字词，词频写入excel文件
#     ws['A1'] = "关键词"
#     ws['B2'] = "词频"
#     for i,row in enumerate(keywords):
#         ws['A{}'.format(i+2)] = row[0]
#         ws['B{}'.format(i+2)] = row[1]
#
#     #保存文件
#     wb.save("统计.xlsx")











